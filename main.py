from typing import List

import numpy as np
import pywintypes
from docx import Document
import pandas as pd
from docx.text.paragraph import Paragraph
from jinja2 import Environment, BaseLoader, TemplateSyntaxError
import argparse
import os
from docx2pdf import convert


class TokenException(BaseException):
    pass


def fill(string, **kwargs):
    try:
        template = Environment(loader=BaseLoader()).from_string(string)
        return template.render(**kwargs)
    except TemplateSyntaxError:
        raise TokenException(f"Wrong token in run: {string}")


def main():
    arg_parser = argparse.ArgumentParser(description="Docx fill app")
    arg_parser.add_argument("docx", type=str, help="Docx file path")
    arg_parser.add_argument("table", type=str, help="Table file path (.xlsx)")
    arg_parser.add_argument("--path", type=str, default="", help="Path where all files wil be stored")
    arg_parser.add_argument("--save_failed", action="store_true", help="To export failed rows")
    arg_parser.add_argument("--render_pdf", action="store_true", help="To render docx to pdf")
    arg_parser.add_argument("--version", action="version", version="Docx fill app 0.2")

    args = arg_parser.parse_args()

    docx_path = args.docx
    table_path = args.table

    table = pd.read_excel(table_path)
    table_copy = None

    if args.save_failed:
        table_copy = table.copy()
    table_dict = table.to_dict()

    idx_to_drop = []

    if not os.path.exists(args.path):
        os.makedirs(args.path, exist_ok=True)

    for idx in range(len(table)):
        docx_file = Document(docx_path)

        def f(pair):
            if pair[1][idx] is np.nan:
                raise ValueError
            return pair[0], pair[1][idx]

        try:
            kwargs = {a: b for a, b in map(f, table_dict.items())}
        except ValueError:
            continue

        table_cells = [docx_file.tables[tb].cell(rw, cl) for tb in range(len(docx_file.tables)) for rw in
                       range(len(docx_file.tables[tb].rows)) for cl in range(len(docx_file.tables[tb].columns))]

        paragraphs: List[Paragraph] = [*docx_file.paragraphs, *[ph for cell in table_cells for ph in cell.paragraphs]]

        for paragraph in paragraphs:
            token = None

            for run_idx in range(len(paragraph.runs)):
                run = paragraph.runs[run_idx]

                if token:
                    token.text = token.text + run.text
                    if "}}" in run.text and "{{" not in run.text or (run.text.startswith("}") and token.text.endswith("}")):
                        token = None
                    run.text = ""

                if "{{" in run.text:
                    try:
                        run.text = fill(run.text, **kwargs)
                    except TokenException:
                        token = run
                elif run.text.endswith("{") and paragraph.runs[run_idx + 1].text.startswith("{"):
                    token = run

        for paragraph in paragraphs:
            for run in paragraph.runs:
                run.text = fill(run.text, **kwargs)

        new_doc_path = fill(os.path.join(args.path, docx_path), **kwargs)
        buff = new_doc_path

        idx = 0
        while os.path.exists(new_doc_path):
            sep = buff.split(".")
            new_doc_path = "".join(sep[:-1]) + f" ({(idx := idx + 1)})." + sep[-1]

        docx_file.save(new_doc_path)
        print("Ready: ", new_doc_path)

        idx_to_drop.append(idx)

    if args.render_pdf:
        try:
            convert(args.path)
        except pywintypes.com_error:
            print("Seems like you dont have Microsoft Word installed.")

    if args.save_failed:
        table_copy = table_copy.drop(idx_to_drop)
    if args.save_failed and len(table_copy):
        table_copy.to_excel("failed_rows.xlsx")


if __name__ == '__main__':
    main()
